import logging
import json
import os
import time
import re
import shutil
import base64
import gc
import sys
from pathlib import Path
from typing import Tuple, Optional, Any, List, Dict

from PIL import Image

from src.rag.document_schema import StructuredDocument, DocumentElement, ElementMetadata, BoundingBox, TableStructure, ImageMetadata
from src.rag.document_builder import DocumentBuilder
from src.rag.ollama_client import OllamaClient

logger = logging.getLogger("pipeline")

class MultimodalExtractor:
    """
    Orchestrates the extraction of both running text (for proofreading)
    and the structured document representation (for future RAG/Assistant features).
    Uses a memory-safe page-streaming architecture to scale to large files (200-500 pages)
    without running out of RAM (avoiding std::bad_alloc errors on CPU systems).
    """

    def __init__(
        self, 
        enable_ocr: bool = False,
        enable_table_extraction: bool = True,
        enable_image_extraction: bool = True
    ) -> None:
        self.enable_ocr = enable_ocr
        self.enable_table_extraction = enable_table_extraction
        self.enable_image_extraction = enable_image_extraction

    def _parse_vlm_output(self, text: str) -> dict:
        """Helper to parse structured sections out of VLM's freeform text response."""
        lines = text.split('\n')
        img_type = "Diagram"
        description = text
        objects = []
        entities = []
        keywords = []
        
        current_section = None
        desc_parts = []
        printed_text_parts = []
        for line in lines:
            line_strip = line.strip()
            if not line_strip:
                continue

            # Simple keyword checks to parse VLM response
            if any(term in line_strip.lower() for term in ["image type", "type of image", "image_type"]):
                current_section = "type"
                parts = line_strip.split(":")
                if len(parts) > 1:
                    img_type = parts[1].strip().strip("*`\"'")
            elif any(term in line_strip.lower() for term in ["semantic description", "overall description", "meaning"]):
                current_section = "desc"
                parts = line_strip.split(":")
                if len(parts) > 1 and parts[1].strip():
                    desc_parts.append(parts[1].strip())
            elif any(term in line_strip.lower() for term in ["detected objects", "objects"]):
                current_section = "objects"
                parts = line_strip.split(":")
                if len(parts) > 1 and parts[1].strip():
                    objects.extend([o.strip().strip("-* ") for o in parts[1].split(",")])
            elif any(term in line_strip.lower() for term in ["entities", "detected entities", "people", "names"]):
                current_section = "entities"
                parts = line_strip.split(":")
                if len(parts) > 1 and parts[1].strip():
                    entities.extend([e.strip().strip("-* ") for e in parts[1].split(",")])
            elif any(term in line_strip.lower() for term in ["printed text", "text printed", "printed/overlaid"]):
                current_section = "printed_text"
                parts = line_strip.split(":")
                if len(parts) > 1 and parts[1].strip():
                    printed_text_parts.append(parts[1].strip())
            elif "keywords" in line_strip.lower():
                current_section = "keywords"
                parts = line_strip.split(":")
                if len(parts) > 1 and parts[1].strip():
                    keywords.extend([k.strip().strip("-* ") for k in parts[1].split(",")])
            else:
                if current_section == "desc":
                    desc_parts.append(line_strip)
                elif current_section == "objects":
                    objects.append(line_strip.strip("-* "))
                elif current_section == "entities":
                    entities.append(line_strip.strip("-* "))
                elif current_section == "printed_text":
                    printed_text_parts.append(line_strip.strip("-* "))
                elif current_section == "keywords":
                    keywords.extend([k.strip().strip("-* ") for k in line_strip.split(",")])

        if desc_parts:
            description = "\n".join(desc_parts)

        # Deduplicate list fields
        objects = sorted(list(set([o for o in objects if len(o) > 1])))
        entities = sorted(list(set([e for e in entities if len(e) > 1])))
        keywords = sorted(list(set([k for k in keywords if len(k) > 1])))

        printed_text = " ".join(printed_text_parts).strip()
        NONE_MARKERS = ("none", "n/a", "na", "-", "blank", "no text", "not visible", "not applicable")
        if printed_text.lower().strip(".") in NONE_MARKERS:
            printed_text = ""

        return {
            "image_type": img_type or "Diagram",
            "semantic_description": description.strip(),
            "objects": objects,
            "entities": entities,
            "keywords": keywords,
            "printed_text": printed_text
        }

    def _update_job_status(
        self,
        doc_id: str,
        stage: str,
        percentage: float,
        current_page: int = 0,
        total_pages: int = 0,
        batch_size: int = 0,
        current_batch: int = 0,
        total_batches: int = 0,
        est_remaining_time: str = "N/A"
    ) -> None:
        """Helper to push granular ingestion stats to the global backend JOBS object."""
        try:
            # Dynamically import backend services to avoid circular dependency in CLI mode
            from backend.services import JOBS, save_job_metadata
            if doc_id in JOBS:
                job = JOBS[doc_id]
                job["current_stage"] = stage
                job["progress_percentage"] = percentage
                job["current_page"] = current_page
                job["total_pages"] = total_pages
                job["batch_size"] = batch_size
                job["current_batch"] = current_batch
                job["total_batches"] = total_batches
                job["estimated_remaining_time"] = est_remaining_time
                job["memory_safe_mode"] = True
                save_job_metadata(doc_id)
        except Exception:
            pass

    def _analyze_document(self, file_path: Path) -> dict:
        """
        Part 1: Document Analyzer
        Inspects the file before parsing to determine page counts, scanned/digital format,
        complexity metrics, and automatically select the optimal streaming strategy.
        """
        total_pages = 1
        is_digital = True
        has_images = False
        has_tables = False
        
        suffix = file_path.suffix.lower()
        if suffix == ".pdf":
            try:
                import fitz
                with fitz.open(file_path) as pdf:
                    total_pages = pdf.page_count
                    sample_pages = min(5, total_pages)
                    total_text_len = 0
                    image_count = 0
                    
                    for i in range(sample_pages):
                        page = pdf[i]
                        text = page.get_text("text") or ""
                        total_text_len += len(text.strip())
                        image_count += len(page.get_images())
                    
                    # Scanned PDF check: low text-to-page ratio
                    avg_text_len = total_text_len / sample_pages
                    is_digital = avg_text_len > 100
                    has_images = image_count > 0
                    
                    # Fast scan for tables
                    tbl_indicator_text = ""
                    for i in range(min(15, total_pages)):
                        tbl_indicator_text += (pdf[i].get_text("text") or "").lower()
                    if any(term in tbl_indicator_text for term in ["table", "columns", "index", "|"]):
                        has_tables = True
            except Exception as e:
                logger.warning(f"DocumentAnalyzer failed to read PDF metrics: {e}")
                is_digital = False
                total_pages = 1
        elif suffix == ".docx":
            try:
                import docx
                doc = docx.Document(str(file_path))
                total_pages = max(1, len(doc.paragraphs) // 15)
                is_digital = True
                has_tables = len(doc.tables) > 0
            except Exception:
                pass
        
        # Decide Batching / Config Strategies
        if suffix in (".docx", ".txt"):
            complexity = "Low (Non-PDF digital)"
            batch_size = total_pages
            enable_ocr = False
            estimated_sec_per_page = 0.05
        elif not is_digital:
            complexity = "High (Scanned)"
            batch_size = 5  # Small batch to prevent memory leaks during OCR
            enable_ocr = True
            estimated_sec_per_page = 4.0
        else:
            enable_ocr = False
            if total_pages > 100:
                complexity = "Medium (Large Digital)"
                batch_size = 15
                estimated_sec_per_page = 0.25
            else:
                complexity = "Low (Standard Digital)"
                batch_size = 25
                estimated_sec_per_page = 0.15
                
        if has_tables:
            estimated_sec_per_page += 0.4
            
        est_total_sec = int(total_pages * estimated_sec_per_page)
        est_mem_mb = 1000 + (250 if enable_ocr else 50) * batch_size
        
        # Format remaining time
        if est_total_sec > 60:
            est_time_str = f"{est_total_sec // 60}m {est_total_sec % 60}s"
        else:
            est_time_str = f"{est_total_sec} seconds"

        return {
            "total_pages": total_pages,
            "is_digital": is_digital,
            "has_images": has_images,
            "has_tables": True,
            "complexity": complexity,
            "batch_size": batch_size,
            "enable_ocr": enable_ocr,
            "enable_table_extraction": True,
            "enable_image_extraction": True,
            "estimated_processing_time": est_time_str,
            "estimated_memory_usage": f"{est_mem_mb} MB",
            "mode": "digital" if is_digital else "scanned"
        }

    def _build_fallback_structured_document_for_batch(
        self, 
        file_path: Path, 
        start_page: int, 
        end_page: int,
        file_name: str,
        file_type: str
    ) -> StructuredDocument:
        """
        Lightweight fallback parser using PyMuPDF to extract text -- AND
        images -- from a batch range when Docling conversion fails for
        those pages.

        Docling failing on a batch must never mean the pictures on those
        pages are silently lost: block_type == 1 ("image/vector graphic")
        blocks are captured here as ImageMetadata entries (bbox + page
        number only, image_path left unset) so the main extraction loop's
        existing "crop directly from source PDF" fallback
        (ImageProcessor.crop_image_from_pdf, triggered whenever a picture's
        PNG doesn't already exist on disk) picks them up, grounds them, and
        persists them through the exact same 05_images pipeline as every
        Docling-detected picture.
        """
        import fitz
        elements = []
        images: Dict[str, ImageMetadata] = {}

        with fitz.open(file_path) as doc:
            for page_idx in range(start_page, end_page):
                if page_idx >= doc.page_count:
                    break
                page = doc[page_idx]
                # Batch-local (1-based within this batch), matching the numbering
                # Docling itself produces when it parses the per-batch temp PDF --
                # NOT the absolute page number in the source document. The caller's
                # "Update batch element IDs and page offsets" pass below adds
                # batch_start to convert every element/image produced here (or by
                # the Docling success path) to an absolute page number exactly
                # once; using an absolute number here would make it get shifted
                # a second time.
                page_num = (page_idx - start_page) + 1

                # Enumerate every embedded raster image on this page via
                # PyMuPDF's image-XObject inspection (page.get_image_info),
                # NOT page.get_text("blocks")'s block_type == 1 classifier --
                # verified against real documents that MuPDF's text-blocks
                # layout engine frequently does not surface embedded
                # pictures as block_type 1 at all (e.g. portrait photos laid
                # out in a card/grid), which would silently lose every
                # picture on any page Docling fails to convert.
                try:
                    image_infos = page.get_image_info(xrefs=True) or []
                except Exception as img_info_err:
                    logger.warning(f"Failed to enumerate images on page {page_num}: {img_info_err}")
                    image_infos = []

                for img_idx, info in enumerate(image_infos):
                    bbox = info.get("bbox")
                    if not bbox or len(bbox) != 4:
                        continue
                    x0, y0, x1, y1 = bbox
                    if x1 <= x0 or y1 <= y0:
                        continue
                    img_ref_id = f"#/pictures/fallback_{page_idx}_{img_idx}"
                    images[img_ref_id] = ImageMetadata(
                        image_id=img_ref_id,
                        page_number=page_num,
                        bbox=BoundingBox(l=x0, t=y0, r=x1, b=y1, coord_origin="TOPLEFT"),
                    )

                # Extract text layout blocks (x0, y0, x1, y1, text, block_no,
                # block_type); block_type == 1 ("image/vector graphic") is
                # skipped here since real pictures are already captured
                # above via get_image_info.
                blocks = page.get_text("blocks") or []

                b_idx = 0
                for b in blocks:
                    if len(b) < 7:
                        continue
                    x0, y0, x1, y1, b_text, b_no, b_type = b[:7]

                    if b_type == 1:
                        continue

                    p_text = (b_text or "").strip()
                    if not p_text or len(p_text) < 2:
                        continue

                    # Filter cover page decorative headers / image titles on Page 1 if short
                    if page_num == 1 and len(p_text) < 15 and not p_text.endswith("."):
                        continue

                    ref_id = f"#/texts/fallback_{page_idx}_{b_idx}"
                    b_idx += 1

                    meta = ElementMetadata(
                        page_number=page_num,
                        # PyMuPDF "blocks" coordinates are already top-left-origin
                        # points (unlike Docling's default BOTTOMLEFT convention),
                        # so tag coord_origin explicitly rather than mislabeling them.
                        bbox=BoundingBox(l=x0, t=y0, r=x1, b=y1, coord_origin="TOPLEFT"),
                        image_id=None,
                        table_id=None,
                        caption_id=None,
                        caption_text=None,
                        level=None,
                        parent_id="body",
                        ocr_text=None
                    )
                    elements.append(DocumentElement(
                        id=ref_id,
                        type="paragraph",
                        text=p_text,
                        metadata=meta,
                        hierarchy_path=[]
                    ))

        return StructuredDocument(
            title=file_name,
            file_name=file_name,
            file_type=file_type,
            page_count=end_page - start_page,
            elements=elements,
            tables={},
            images=images
        )

    def extract(
        self, 
        file_path: Path, 
        output_dir: Optional[Path] = None
    ) -> Tuple[str, StructuredDocument, int]:
        """
        Main extraction entry point. Automatically analyzes PDF metrics, configures layout options,
        and streams parsing, knowledge extraction, and vector index updates in page-range batches.
        """
        start_time = time.time()
        suffix = file_path.suffix.lower()
        file_name = file_path.stem
        file_type = suffix.lstrip(".")
        doc_id = output_dir.name if output_dir else file_name

        # Caching optimization check
        if output_dir:
            final_doc_path = output_dir / "02_docling" / "structured_document.json"
            final_ko_path = output_dir / "03_knowledge_objects" / "knowledge_objects.json"
            if final_doc_path.exists() and final_ko_path.exists():
                logger.info("Complete extraction cache hit! Reusing StructuredDocument and Knowledge Objects.")
                with open(final_doc_path, "r", encoding="utf-8") as f:
                    master_dict = json.load(f)
                    master_structured_doc = StructuredDocument(**master_dict)
                with open(final_ko_path, "r", encoding="utf-8") as f:
                    ko_list = json.load(f)
                    total_objs = len(ko_list)
                    
                total_pages = master_structured_doc.page_count
                
                # Ensure legacy structured_document.json exists at root of output_dir
                legacy_doc_path = output_dir / "structured_document.json"
                if not legacy_doc_path.exists():
                    try:
                        with open(legacy_doc_path, "w", encoding="utf-8") as f:
                            json.dump(master_dict, f, indent=2, ensure_ascii=False)
                    except Exception as save_legacy_err:
                        logger.error(f"Failed to save legacy structured document copy: {save_legacy_err}")
                
                # Check ChromaDB indexing status and rebuild index from cache if missing
                try:
                    from src.rag.config import RagConfig
                    from src.rag.vector_store import VectorStore
                    from src.rag.index_manager import IndexManager
                    
                    config = RagConfig()
                    vector_store = VectorStore(
                        db_dir=config.chroma_db_dir,
                        collection_prefix=config.collection_prefix
                    )
                    existing_ids = vector_store.get_existing_chunk_ids(doc_id)
                    if not existing_ids:
                        logger.info("ChromaDB index not found for cached document. Re-indexing from file...")
                        chunks_file = output_dir / "06_chunks" / "document_chunks.json"
                        if not chunks_file.exists():
                            chunks_file = output_dir / "document_chunks.json"
                        if chunks_file.exists():
                            index_manager = IndexManager.from_config(config)
                            index_manager.index_from_file(chunks_file, doc_id)
                except Exception as db_err:
                    logger.error(f"Failed to check/rebuild vector index for cached document: {db_err}")

                # Update global JOBS stats if running inside backend
                try:
                    from backend.services import JOBS, save_job_metadata
                    if doc_id in JOBS:
                        job = JOBS[doc_id]
                        job["knowledge_objects_generated"] = total_objs
                        job["embeddings_completed"] = total_objs
                        job["index_progress"] = 100
                        job["current_page"] = total_pages
                        job["total_pages"] = total_pages
                        job["current_batch"] = 1
                        job["total_batches"] = 1
                        job["estimated_remaining_time"] = "0s"
                        save_job_metadata(doc_id)
                except Exception:
                    pass
                    
                # Compile final raw text
                raw_text = master_structured_doc.title + "\n\n" + "\n\n".join(
                    el.text for el in master_structured_doc.elements if el.type in ("paragraph", "heading", "list_item")
                )
                self._update_job_status(doc_id, "Completed Ingestion (Cached)", 100.0, total_pages, total_pages, 1, 1, 1, "0s")
                return raw_text, master_structured_doc, total_pages

        # Create output workspaces
        if output_dir:
            for folder in ["01_document", "02_docling", "03_knowledge_objects", "04_tables", "05_images", "06_embeddings", "07_vectordb", "08_retrieval", "09_reports"]:
                (output_dir / folder).mkdir(parents=True, exist_ok=True)
            # Copy uploaded original
            shutil.copy2(file_path, output_dir / "01_document" / f"uploaded{suffix}")

        # Part 1: Run Document Analyzer
        self._update_job_status(doc_id, "Analyzing Document", 2.0)
        analysis = self._analyze_document(file_path)
        logger.info(f"Document Ingestion Strategy: {json.dumps(analysis, indent=2)}")

        total_pages = analysis["total_pages"]
        batch_size = analysis["batch_size"]
        self.enable_ocr = analysis["enable_ocr"]
        self.enable_table_extraction = analysis["enable_table_extraction"]
        self.enable_image_extraction = analysis["enable_image_extraction"]

        # Part 5: Checkpointing - Load existing progress if crashed
        checkpoint_path = output_dir / "ingestion_checkpoint.json" if output_dir else None
        master_structured_doc = None
        start_page = 0
        skipped_pages = []
        cumulative_stats = {}

        if checkpoint_path and checkpoint_path.exists():
            try:
                with open(checkpoint_path, "r", encoding="utf-8") as f:
                    checkpoint = json.load(f)
                start_page = checkpoint.get("current_page", 0)
                skipped_pages = checkpoint.get("skipped_pages", [])
                cumulative_stats = checkpoint.get("stats", {})
                logger.info(f"Ingestion checkpoint detected. Resuming from page {start_page}...")

                # Reload master structured document
                master_doc_path = output_dir / "02_docling" / "structured_document.json"
                if master_doc_path.exists():
                    with open(master_doc_path, "r", encoding="utf-8") as f:
                        master_structured_doc = StructuredDocument(**json.load(f))
            except Exception as checkpoint_err:
                logger.error(f"Failed to load checkpoint: {checkpoint_err}. Restarting ingestion.")
                start_page = 0
                skipped_pages = []

        if not master_structured_doc:
            master_structured_doc = StructuredDocument(
                title=file_name,
                file_name=file_name,
                file_type=file_type,
                page_count=total_pages,
                elements=[],
                tables={},
                images={}
            )

        # Initialize Knowledge Extraction Agent
        from src.rag.knowledge_extraction_agent import KnowledgeExtractionAgent
        from src.rag.config import RagConfig
        agent = KnowledgeExtractionAgent(config=RagConfig())
        if cumulative_stats:
            agent.stats.update(cumulative_stats)

        # Content-based image deduplication registry, persisted across the
        # whole streaming extraction so a logo/header/footer repeated in a
        # later batch is still caught against canonicals from earlier batches.
        from src.rag.image_deduplicator import ImageDeduplicationRegistry
        dedup_registry = ImageDeduplicationRegistry()
        if output_dir and start_page > 0:
            # Resuming from a checkpoint: prime the registry with images already
            # written by earlier (pre-crash) batches so duplicates spanning the
            # restart boundary are still caught.
            existing_images_dir = output_dir / "05_images"
            if existing_images_dir.exists():
                for existing_png in sorted(existing_images_dir.glob("image_*.png")):
                    dedup_registry.check_and_register(existing_png, existing_png.stem)

        # Initialize Docling Converter once
        converter = None
        try:
            from docling.document_converter import DocumentConverter, PdfFormatOption
            from docling.datamodel.pipeline_options import PdfPipelineOptions
            from docling.datamodel.base_models import InputFormat

            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = self.enable_ocr
            pipeline_options.do_table_structure = self.enable_table_extraction
            pipeline_options.generate_picture_images = self.enable_image_extraction
            try:
                # Docling defaults to only 4 CPU threads regardless of host
                # size; use all available cores unless explicitly overridden.
                from docling.datamodel.pipeline_options import AcceleratorOptions
                pipeline_options.accelerator_options = AcceleratorOptions(
                    num_threads=int(os.environ.get("DOCLING_NUM_THREADS", str(os.cpu_count() or 4))),
                    device=os.environ.get("DOCLING_DEVICE", "auto"),
                )
            except ImportError:
                pass
            try:
                from docling.datamodel.pipeline_options import RapidOcrOptions
                pipeline_options.ocr_options = RapidOcrOptions()
            except ImportError:
                pass
            
            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                }
            )
        except ImportError:
            logger.warning("Docling package is not installed.")

        # Part 2: Page Batch Loop
        total_batches = (total_pages + batch_size - 1) // batch_size
        current_page = start_page
        batch_index = current_page // batch_size

        while current_page < total_pages:
            batch_start = current_page
            batch_end = min(batch_start + batch_size, total_pages)
            batch_index += 1
            
            # Progress Tracking
            percent = 5.0 + (batch_start / total_pages) * 75.0
            elapsed = time.time() - start_time
            pages_processed = batch_start - start_page
            pages_left = total_pages - batch_start
            
            if pages_processed > 0:
                est_seconds_left = int((elapsed / pages_processed) * pages_left)
                if est_seconds_left > 60:
                    est_time_str = f"{est_seconds_left // 60}m {est_seconds_left % 60}s"
                else:
                    est_time_str = f"{est_seconds_left}s"
            else:
                est_time_str = "Estimating..."

            self._update_job_status(
                doc_id=doc_id,
                stage="Extracting Layout",
                percentage=percent,
                current_page=batch_start + 1,
                total_pages=total_pages,
                batch_size=batch_size,
                current_batch=batch_index,
                total_batches=total_batches,
                est_remaining_time=est_time_str
            )
            
            logger.info(f"Processing page batch {batch_index}/{total_batches} (Pages {batch_start+1} - {batch_end} of {total_pages})...")

            batch_structured_doc = None
            docling_doc = None
            
            # Extract PDF batch range to temp file
            temp_pdf_path = output_dir / f"temp_batch_{batch_start}_{batch_end}.pdf" if output_dir else Path(f"temp_batch.pdf")
            
            try:
                import fitz
                with fitz.open(file_path) as pdf:
                    batch_pdf = fitz.open()
                    batch_pdf.insert_pdf(pdf, from_page=batch_start, to_page=batch_end - 1)
                    batch_pdf.save(str(temp_pdf_path))
                    batch_pdf.close()
            except Exception as extract_err:
                logger.error(f"Failed to crop batch pages: {extract_err}")
                temp_pdf_path = file_path # Fallback

            # Convert page range via Docling
            # Part 8: Fault Tolerance with lightweight PyMuPDF Fallback
            docling_success = False
            for retry in range(2):
                try:
                    docling_doc, raw_text, page_count = self._run_docling_conversion(converter, temp_pdf_path, output_dir)
                    if docling_doc is not None:
                        images_dir = output_dir / "05_images" if output_dir else None
                        if images_dir:
                            images_dir.mkdir(parents=True, exist_ok=True)
                        builder = DocumentBuilder(
                            output_images_dir=images_dir,
                            pdf_path=file_path
                        )
                        batch_structured_doc = builder.build(docling_doc, file_name, file_type)
                        docling_success = True
                        break
                except Exception as docling_err:
                    logger.warning(f"Docling batch conversion failed (Attempt {retry+1}): {docling_err}")
                    time.sleep(1)

            if not docling_success:
                logger.error(f"Docling batch conversion failed after retries. Invoking lightweight fallback for pages {batch_start+1}-{batch_end}.")
                skipped_pages.extend(range(batch_start + 1, batch_end + 1))
                batch_structured_doc = self._build_fallback_structured_document_for_batch(
                    file_path=file_path,
                    start_page=batch_start,
                    end_page=batch_end,
                    file_name=file_name,
                    file_type=file_type
                )

            # Cleanup temp batch PDF
            if temp_pdf_path.exists() and temp_pdf_path != file_path:
                try:
                    temp_pdf_path.unlink()
                except Exception:
                    pass

            # Update batch element IDs and page offsets to align with the absolute master document
            for el in batch_structured_doc.elements:
                el.id = f"b{batch_index}_{el.id}"
                if el.metadata.parent_id and el.metadata.parent_id != "body":
                    el.metadata.parent_id = f"b{batch_index}_{el.metadata.parent_id}"
                el.hierarchy_path = [f"b{batch_index}_{pid}" for pid in el.hierarchy_path]
                el.metadata.page_number = el.metadata.page_number + batch_start
                if el.metadata.table_id:
                    el.metadata.table_id = f"b{batch_index}_{el.metadata.table_id}"
                if el.metadata.image_id:
                    el.metadata.image_id = f"b{batch_index}_{el.metadata.image_id}"

            # Images carry their own page_number (Docling assigns it from the
            # per-batch temp PDF it parsed, so it is batch-local -- 1..batch_size
            # -- exactly like the elements above before the offset was applied).
            # It was never being shifted to the absolute page number here, which
            # meant every image after the first batch was persisted with the
            # wrong "page" in its JSON, and its caption/layout grounding lookup
            # below (which matches against elements' *already-absolute*
            # page_number) silently found zero surrounding page elements for
            # every batch past the first -- degrading grounding/retrievability
            # for the majority of a large document's images without dropping
            # the PNG itself, and pointing any PDF-crop fallback at the wrong
            # page of the source document.
            for img_meta in batch_structured_doc.images.values():
                img_meta.page_number = img_meta.page_number + batch_start

            # Merge into master document layout tree
            master_structured_doc.elements.extend(batch_structured_doc.elements)
            for k, val in batch_structured_doc.tables.items():
                master_structured_doc.tables[f"b{batch_index}_{k}"] = val
            for k, val in batch_structured_doc.images.items():
                master_structured_doc.images[f"b{batch_index}_{k}"] = val

            # Persist master layout on disk
            if output_dir:
                try:
                    structured_doc_path = output_dir / "02_docling" / "structured_document.json"
                    structured_doc_path.parent.mkdir(parents=True, exist_ok=True)
                    with open(structured_doc_path, "w", encoding="utf-8") as f:
                        json.dump(master_structured_doc.model_dump(), f, indent=2, ensure_ascii=False)
                except Exception as save_err:
                    logger.error(f"Failed to save intermediate structured document: {save_err}")

            # Run VLM processing for images extracted in this batch
            if output_dir and batch_structured_doc.images:
                self._update_job_status(doc_id, "Generating Knowledge Objects", percent, batch_start+1, total_pages, batch_size, batch_index, total_batches, est_time_str)
                ollama_client = agent.ollama_client
                for idx, (img_id, img_meta) in enumerate(list(batch_structured_doc.images.items())):
                    orig_path = Path(img_meta.image_path) if img_meta.image_path else None
                    seq_name = f"image_{len(master_structured_doc.images) - len(batch_structured_doc.images) + idx + 1:03d}"
                    new_png_path = output_dir / "05_images" / f"{seq_name}.png"
                    
                    if orig_path and orig_path.exists():
                        if orig_path.resolve() != new_png_path.resolve():
                            try:
                                if new_png_path.exists():
                                    new_png_path.unlink()
                                shutil.move(str(orig_path), str(new_png_path))
                                img_meta.image_path = str(new_png_path)
                                master_structured_doc.images[f"b{batch_index}_{img_id}"].image_path = str(new_png_path)
                            except Exception as move_err:
                                logger.error(f"Failed to sequentialize image {seq_name}: {move_err}")
                        else:
                            img_meta.image_path = str(new_png_path)
                            master_structured_doc.images[f"b{batch_index}_{img_id}"].image_path = str(new_png_path)

                    # Fallback extraction: If new_png_path still does not exist or is empty, crop directly from source PDF
                    if (not new_png_path.exists() or new_png_path.stat().st_size == 0) and file_path and img_meta.bbox:
                        ImageProcessor.crop_image_from_pdf(
                            pdf_path=Path(file_path),
                            page_number=img_meta.page_number,
                            bbox=img_meta.bbox,
                            target_path=new_png_path
                        )
                        if new_png_path.exists() and new_png_path.stat().st_size > 0:
                            img_meta.image_path = str(new_png_path)
                            if f"b{batch_index}_{img_id}" in master_structured_doc.images:
                                master_structured_doc.images[f"b{batch_index}_{img_id}"].image_path = str(new_png_path)
                    
                    # Strictly ensure physical PNG existence before continuing
                    if not new_png_path.exists() or new_png_path.stat().st_size == 0:
                        logger.warning(f"Physical image file could not be created for {seq_name} (Page {img_meta.page_number}). Skipping metadata save to enforce 1:1 mapping.")
                        continue

                    # Content-based duplicate detection (exact + near-duplicate
                    # perceptual hash), BEFORE any grounding, VLM captioning,
                    # JSON metadata write, or chunking/indexing happens for
                    # this image. Repeated logos, headers, footers, decorative
                    # assets, and any other pixel-identical/near-identical
                    # repeat are collapsed onto their earlier canonical copy;
                    # genuinely different images are never touched.
                    canonical_seq = dedup_registry.check_and_register(new_png_path, seq_name)
                    if canonical_seq is not None:
                        from src.rag.image_deduplicator import merge_duplicate_page_into_canonical
                        canonical_json_path = output_dir / "05_images" / f"{canonical_seq}.json"
                        merge_duplicate_page_into_canonical(
                            canonical_json_path=canonical_json_path,
                            duplicate_page_number=img_meta.page_number,
                            duplicate_image_id=img_meta.image_id,
                        )
                        try:
                            new_png_path.unlink()
                        except Exception:
                            pass
                        # Remove from every structure chunk_builder / master doc
                        # will read, so the duplicate is invisible to chunking
                        # and indexing (not merely undescribed).
                        master_structured_doc.images.pop(f"b{batch_index}_{img_id}", None)
                        batch_structured_doc.images.pop(img_id, None)
                        batch_structured_doc.elements = [
                            el for el in batch_structured_doc.elements
                            if not (
                                el.type == "image"
                                and (el.id.endswith(img_id) or (el.metadata.image_id or "").endswith(img_id))
                            )
                        ]
                        logger.info(
                            f"Duplicate image detected: {seq_name} (Page {img_meta.page_number}) "
                            f"merged into canonical {canonical_seq}; excluded from chunking/indexing."
                        )
                        continue

                    # Step 3: Hierarchical Caption Detection & Layout Association
                    page_elements = [
                        el for el in batch_structured_doc.elements 
                        if getattr(el.metadata, "page_number", None) == img_meta.page_number
                    ]
                    page_elements_dicts = []
                    for pe in page_elements:
                        b_dict = pe.metadata.bbox.model_dump() if hasattr(pe.metadata.bbox, "model_dump") else (pe.metadata.bbox.dict() if pe.metadata.bbox else None)
                        page_elements_dicts.append({
                            "id": pe.id,
                            "type": pe.type,
                            "text": pe.text,
                            "metadata": {"bbox": b_dict, "page_number": pe.metadata.page_number}
                        })

                    from src.rag.image_processor import HierarchicalLayoutGrounder
                    grounded = HierarchicalLayoutGrounder.ground_image(
                        image_id=img_meta.image_id,
                        page_number=img_meta.page_number,
                        bbox=img_meta.bbox,
                        doc_elements_on_page=page_elements_dicts,
                        doc_title=batch_structured_doc.title,
                        active_section=img_meta.related_section,
                        explicit_caption=img_meta.caption,
                        ocr_text=img_meta.ocr_text,
                        raw_image_type=img_meta.image_type,
                        vlm_description=img_meta.semantic_description
                    )

                    # Update img_meta with grounded attributes
                    img_meta.title = grounded.get("title")
                    img_meta.subtitle = grounded.get("subtitle")
                    img_meta.explicit_caption = grounded["explicit_caption"]
                    img_meta.caption_text = grounded["caption_text"]
                    img_meta.caption = grounded["caption"]
                    img_meta.entity_name = grounded["entity_name"]
                    img_meta.designation = grounded["designation"]
                    img_meta.section_heading = grounded["section_heading"]
                    img_meta.text_before = grounded["text_before"]
                    img_meta.text_after = grounded["text_after"]
                    img_meta.nearby_text = grounded.get("nearby_text")
                    img_meta.layout_context = grounded["layout_context"]
                    img_meta.image_type = grounded["image_type"]
                    img_meta.importance_score = grounded["importance_score"]
                    img_meta.retrievable = grounded["retrievable"]
                    img_meta.association_method = grounded["association_method"]
                    img_meta.association_confidence = grounded["association_confidence"]
                    img_meta.confidence = grounded["confidence"]

                    # VLM enrichment for meaningful visual assets (never skipped prematurely)
                    if new_png_path.exists() and not img_meta.semantic_description:
                        try:
                            with open(new_png_path, "rb") as image_file:
                                img_b64 = base64.b64encode(image_file.read()).decode("utf-8")
                            
                            prompt = (
                                f"Describe this image extracted from a document.\n"
                                f"Caption: {img_meta.caption or ''}\n"
                                f"OCR text: {img_meta.ocr_text or ''}\n"
                                "Please provide a structured response in natural language describing:\n"
                                "1. Image Type (e.g. Graph, Chart, Map, Photo, Diagram, Flowchart, Logo, Screenshot, Portrait Photo)\n"
                                "2. Semantic Description (a detailed explanation of the overall meaning)\n"
                                "3. Detected Objects (list of main objects/elements)\n"
                                "4. Detected Entities (people, names, or organizations if any)\n"
                                "5. Keywords (a list of comma-separated keywords)\n"
                                "6. Printed Text (transcribe verbatim any name, caption, designation, or other text\n"
                                "   actually printed/overlaid on the image itself, e.g. a name label under a portrait;\n"
                                "   leave blank if none is visible)\n"
                                "\nAnswer clearly and directly."
                            )
                            vlm_text = ollama_client.generate_vision(
                                model="qwen2.5vl:latest",
                                prompt=prompt,
                                image_bytes_b64=img_b64
                            )
                            parsed = self._parse_vlm_output(vlm_text)
                            # A grounded image_type is only trusted here when it came from a
                            # real document signal (an explicit caption, a same-card name/photo
                            # layout, or OCR-read identity). Anything else -- including Docling's
                            # own built-in picture classifier, which can mislabel a full-page
                            # portrait as e.g. "Chart/Graph" with no caption to cross-check against --
                            # is a low-confidence guess the VLM's actual look at the pixels may correct.
                            if img_meta.association_method not in ("same_card_layout", "explicit_caption", "ocr_grounded_identity"):
                                img_meta.image_type = parsed.get("image_type", img_meta.image_type)
                            img_meta.semantic_description = parsed.get("semantic_description", "")
                            img_meta.objects = parsed.get("objects", [])
                            img_meta.detected_entities = parsed.get("entities", [])
                            img_meta.keywords = parsed.get("keywords", [])
                            if not img_meta.ocr_text and parsed.get("printed_text"):
                                img_meta.ocr_text = parsed["printed_text"]
                        except Exception as vlm_err:
                            logger.warning(f"VLM analysis failed for {seq_name}: {vlm_err}")

                    if not img_meta.semantic_description:
                        img_meta.semantic_description = grounded.get("semantic_description") or f"Document visual graphic on Page {img_meta.page_number}."

                    # Late signature-grounded identity promotion: checked
                    # BEFORE the generic OCR-grounded portrait promotion
                    # below, since a signature block's OCR text also
                    # contains a titled name and would otherwise always be
                    # misread as a portrait. Uses only generic English
                    # signature phrases, never any person/organization name.
                    is_late_signature_block = False
                    if not img_meta.entity_name and img_meta.ocr_text:
                        _SIG_INDICATORS = (
                            "signature", "signed by", "authorised signatory",
                            "authorized signatory", "for and on behalf of", "sd/-"
                        )
                        if any(ind in img_meta.ocr_text.lower() for ind in _SIG_INDICATORS):
                            is_late_signature_block = True
                            from src.rag.image_processor import extract_generic_person_identity, extract_untitled_name_near_designation
                            sig_identity = extract_generic_person_identity(img_meta.ocr_text) or extract_untitled_name_near_designation(img_meta.ocr_text)
                            if sig_identity:
                                entity_name, designation = sig_identity
                                img_meta.entity_name = entity_name
                                img_meta.designation = designation
                                img_meta.title = entity_name
                                img_meta.subtitle = designation or "Signature"
                                img_meta.image_type = "Signature"
                                img_meta.layout_context = "signature_block"
                                img_meta.association_method = "signature_text_grounded"
                                img_meta.association_confidence = 0.85
                                img_meta.confidence = 0.85
                                img_meta.caption_text = f"Signature of {entity_name} ({designation})" if designation else f"Signature of {entity_name}"
                                img_meta.caption = img_meta.caption_text
                                img_meta.explicit_caption = img_meta.caption_text

                    # Late OCR-grounded identity promotion: the layout grounder runs before
                    # the VLM call above and can only see caption/OCR text known at that time.
                    # If the VLM's transcription of text printed on the image itself (or any
                    # OCR text otherwise present) reveals a name/designation that grounding
                    # missed, apply the same real-text-only identity rule here so images whose
                    # only identifying signal is baked into the pixels (no separate caption
                    # text-block for Docling to parse) still get correctly associated.
                    if not img_meta.entity_name and not is_late_signature_block and img_meta.ocr_text:
                        from src.rag.image_processor import extract_generic_person_identity, extract_untitled_name_near_designation
                        ocr_identity = extract_generic_person_identity(img_meta.ocr_text) or extract_untitled_name_near_designation(img_meta.ocr_text)
                        if ocr_identity:
                            entity_name, designation = ocr_identity
                            img_meta.entity_name = entity_name
                            img_meta.designation = designation
                            img_meta.title = entity_name
                            img_meta.subtitle = designation
                            img_meta.image_type = "Portrait Photo"
                            img_meta.layout_context = "ocr_grounded_portrait"
                            img_meta.association_method = "ocr_grounded_identity"
                            img_meta.association_confidence = 0.80
                            img_meta.confidence = 0.80
                            img_meta.caption_text = f"Portrait of {entity_name} ({designation})" if designation else f"Portrait of {entity_name}"
                            img_meta.caption = img_meta.caption_text
                            img_meta.explicit_caption = img_meta.caption_text

                    # Post-VLM Retrievability & Score Re-evaluation
                    is_logo_semantic = (
                        img_meta.image_type == "Logo" or
                        "logo" in (img_meta.semantic_description or "").lower() or
                        any("logo" in str(k).lower() for k in (img_meta.keywords or []))
                    )
                    if is_logo_semantic:
                        from src.rag.image_processor import extract_generic_organization_name
                        org_name = (
                            extract_generic_organization_name([
                                img_meta.explicit_caption, img_meta.semantic_description,
                                img_meta.ocr_text, batch_structured_doc.title
                            ])
                            or (batch_structured_doc.title.strip() if batch_structured_doc.title else None)
                            or "the associated organization"
                        )
                        img_meta.image_type = "Logo"
                        img_meta.importance_score = "HIGH"
                        img_meta.retrievable = True
                        img_meta.title = f"Company Logo - {org_name}"
                        img_meta.subtitle = "Official Brand Identity"
                        img_meta.caption_text = f"Company Logo of {org_name}"
                        img_meta.caption = img_meta.caption_text
                        img_meta.keywords = list(set((img_meta.keywords or []) + ["logo", "company logo", "brand", "emblem", "insignia", org_name]))
                        img_meta.detected_entities = list(set((img_meta.detected_entities or []) + [org_name]))
                    elif img_meta.image_type in ("Portrait", "Portrait Photo") or img_meta.entity_name:
                        img_meta.importance_score = "HIGH"
                        img_meta.retrievable = True
                    elif img_meta.image_type in ("Chart", "Graph", "Chart/Graph", "Diagram", "Table", "Map", "Signature"):
                        img_meta.importance_score = "HIGH"
                        img_meta.retrievable = True

                    # Save image meta JSON strictly conforming to 20-field unified schema
                    image_json_path = output_dir / "05_images" / f"{seq_name}.json"
                    bbox_data = img_meta.bbox.model_dump() if hasattr(img_meta.bbox, "model_dump") else (img_meta.bbox.dict() if img_meta.bbox else None)
                    img_json_data = {
                        "image_id": img_meta.image_id,
                        "image_path": f"05_images/{seq_name}.png",
                        "image_url": f"/outputs/{doc_id}/05_images/{seq_name}.png",
                        "page": img_meta.page_number,
                        "bounding_box": bbox_data,
                        "image_type": img_meta.image_type,
                        "title": img_meta.title or f"Visual on Page {img_meta.page_number}",
                        "subtitle": img_meta.subtitle,
                        "explicit_caption": img_meta.explicit_caption,
                        "caption_text": img_meta.caption_text,
                        "entity_name": img_meta.entity_name,
                        "designation": img_meta.designation,
                        "section_heading": img_meta.section_heading,
                        "text_before": img_meta.text_before,
                        "text_after": img_meta.text_after,
                        "nearby_text": img_meta.nearby_text,
                        "semantic_description": img_meta.semantic_description,
                        "keywords": img_meta.keywords or grounded.get("keywords", []),
                        "importance_score": img_meta.importance_score,
                        "retrievable": img_meta.retrievable,
                        "association_method": img_meta.association_method,
                        "association_confidence": img_meta.association_confidence,
                        # Backward-compatible fields
                        "caption": img_meta.caption,
                        "confidence": img_meta.confidence,
                        "layout_context": img_meta.layout_context,
                        "ocr_text": img_meta.ocr_text,
                        "objects": img_meta.objects or grounded.get("objects", []),
                        "detected_entities": img_meta.detected_entities or grounded.get("detected_entities", [])
                    }
                    with open(image_json_path, "w", encoding="utf-8") as f:
                        json.dump(img_json_data, f, indent=2, ensure_ascii=False)

            # Part 3 & 4: Incremental Knowledge Extraction & Vector DB Updates (Semantic Chunks reasoning unit)
            self._update_job_status(doc_id, "Building Embeddings", percent, batch_start+1, total_pages, batch_size, batch_index, total_batches, est_time_str)
            batch_objects = []
            for obj in agent._generate_objects(batch_structured_doc, doc_id, output_dir):
                batch_objects.append(obj)
                t = obj.chunk_type
                agent.stats["by_type"][t] = agent.stats["by_type"].get(t, 0) + 1
                agent.stats["total_objects"] += 1
            
            if batch_objects and output_dir:
                jsonl_path = output_dir / "03_knowledge_objects" / "knowledge_objects.jsonl"
                jsonl_path.parent.mkdir(parents=True, exist_ok=True)
                with open(jsonl_path, "a", encoding="utf-8") as f:
                    for obj in batch_objects:
                        obj.embedding_status = "Pending"
                        obj_dict = obj.model_dump() if hasattr(obj, "model_dump") else obj.dict()
                        f.write(json.dumps(obj_dict, ensure_ascii=False) + "\n")
            
            # Rebuilt Semantic Chunking (Tasks 1-8)
            from src.rag.chunk_builder import ChunkBuilder
            overlap_size = getattr(agent.config, "chunk_overlap_tokens", 50)
            chunk_builder = ChunkBuilder(
                target_tokens_min=250, 
                target_tokens_max=500, 
                overlap_tokens=overlap_size
            )
            
            batch_chunks = chunk_builder.build_chunks(batch_structured_doc, doc_id)
            if batch_chunks:
                agent._process_chunks_batch(batch_chunks, doc_id, output_dir)

            # Part 6: Memory Management - Release layout objects and tensors
            del docling_doc
            del batch_structured_doc
            del batch_objects
            gc.collect()

            try:
                import torch
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
            except ImportError:
                pass

            # Save check-point state
            current_page = batch_end
            if checkpoint_path:
                try:
                    with open(checkpoint_path, "w", encoding="utf-8") as f:
                        json.dump({
                            "current_page": current_page,
                            "skipped_pages": skipped_pages,
                            "stats": agent.stats
                        }, f, indent=2)
                except Exception as cp_save_err:
                    logger.error(f"Failed to write ingestion checkpoint: {cp_save_err}")

        # Post-Processing: Finalize Knowledge objects and chunks
        self._update_job_status(doc_id, "Updating Vector Database", 82.0, total_pages, total_pages, batch_size, total_batches, total_batches, "0s")
        # Make dedup stats available to _finalize_outputs' image artifact
        # cleanup/validation pass before it runs, so its funnel log is accurate.
        agent.stats["image_dedup"] = dedup_registry.summary()
        if output_dir:
            agent._finalize_outputs(doc_id, output_dir, master_structured_doc)

            # Save structured_document.json copy at root of output_dir for compatibility
            try:
                legacy_doc_path = output_dir / "structured_document.json"
                if hasattr(master_structured_doc, "model_dump_json"):
                    json_data = master_structured_doc.model_dump_json(indent=2)
                else:
                    json_data = master_structured_doc.json(indent=2)
                legacy_doc_path.write_text(json_data, encoding="utf-8")
                logger.info(f"Saved structured_document.json copy to {legacy_doc_path}")
            except Exception as save_err:
                logger.error(f"Failed to save structured_document.json copy: {save_err}")

        # Finalize knowledge objects outputs (JSON, HTML, Markdown)
        if output_dir:
            agent._finalize_outputs(doc_id, output_dir, master_structured_doc)

            # Enforce canonical 05_images persistence: remove any stray staging or non-canonical files
            images_dir = output_dir / "05_images"
            if images_dir.exists():
                for f in images_dir.iterdir():
                    if f.is_file() and not f.name.startswith("image_"):
                        try:
                            f.unlink()
                            logger.info(f"Removed non-canonical file from 05_images: {f.name}")
                        except Exception as e:
                            logger.warning(f"Could not remove non-canonical file {f.name}: {e}")

        # Image deduplication summary: total extracted -> duplicates removed -> unique retained -> images indexed
        dedup_stats = dedup_registry.summary()
        images_indexed = agent.stats.get("by_type", {}).get("image", 0)
        logger.info(
            "Image pipeline summary | total extracted: %d -> duplicates removed: %d -> "
            "unique retained: %d -> images indexed: %d",
            dedup_stats["total_extracted"],
            dedup_stats["duplicates_removed"],
            dedup_stats["unique_retained"],
            images_indexed,
        )

        # Ingestion Completed
        self._update_job_status(doc_id, "Ingestion Completed", 100.0, total_pages, total_pages, batch_size, total_batches, total_batches, "0s")

        # Clean checkpoint file upon success
        if checkpoint_path and checkpoint_path.exists():
            try:
                checkpoint_path.unlink()
            except Exception:
                pass

        # Write readme logs
        processing_time = time.time() - start_time
        if output_dir:
            vlm_count = agent.stats.get("vision_processed_images", 0)
            total_objs = agent.stats.get("total_objects", 0)
            readme_text = (
                f"# Multimodal Knowledge Ingestion Summary\n\n"
                f"## Statistics\n"
                f"- **Job ID**: `{doc_id}`\n"
                f"- **Pages Parsed**: {total_pages}\n"
                f"- **Skipped Pages (Failures)**: {len(skipped_pages)} ({json.dumps(skipped_pages)})\n"
                f"- **Knowledge Objects Created**: {total_objs}\n"
                f"- **VLM Vision Descriptions**: {vlm_count}\n"
                f"- **Processing Strategy**: {analysis['complexity']} (Safe Streaming Mode)\n"
                f"- **Total Processing Time**: {processing_time:.2f} seconds\n\n"
                f"## Image Deduplication\n"
                f"- **Total Images Extracted**: {dedup_stats['total_extracted']}\n"
                f"- **Duplicate Images Removed**: {dedup_stats['duplicates_removed']}\n"
                f"- **Unique Images Retained**: {dedup_stats['unique_retained']}\n"
                f"- **Images Indexed**: {images_indexed}\n\n"
                f"## Collections\n"
                f"- **ChromaDB**: `doc_{doc_id}`\n"
            )
            (output_dir / "README.md").write_text(readme_text, encoding="utf-8")

        # Compile final raw text from layout elements for downstream proofreading stages
        raw_text = master_structured_doc.title + "\n\n" + "\n\n".join(
            el.text for el in master_structured_doc.elements if el.type in ("paragraph", "heading", "list_item")
        )

        if converter:
            del converter

        return raw_text, master_structured_doc, total_pages

    def _run_docling_conversion(
        self, 
        converter: Any,
        file_path: Path, 
        output_dir: Optional[Path] = None
    ) -> Tuple[Optional[Any], Optional[str], int]:
        """Runs Docling converter over the temp batch document."""
        if converter is None:
            return None, None, 0
        try:
            result = converter.convert(str(file_path))
            docling_doc = result.document
            
            markdown = docling_doc.export_to_markdown()
            raw_text = self._markdown_to_text(markdown)
            page_count = len(getattr(docling_doc, "pages", []) or []) or 1
            return docling_doc, raw_text, page_count
        except Exception as e:
            logger.error(f"Docling conversion error: {e}")
            return None, None, 0

    def _run_fallback_conversion(self, file_path: Path, suffix: str) -> Tuple[str, int]:
        if suffix == ".pdf":
            try:
                import fitz
                text_parts = []
                with fitz.open(file_path) as pdf:
                    page_count = pdf.page_count
                    for page in pdf:
                        text_parts.append(page.get_text("text"))
                return "\n\n".join(text_parts), page_count
            except Exception as e:
                logger.error(f"Fallback PDF extractor failed: {e}")
        return "", 0

    @staticmethod
    def _markdown_to_text(markdown: str) -> str:
        from src.utils import strip_markdown
        return strip_markdown(markdown)
