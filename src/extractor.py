"""
extractor.py
============
Stage 1 of the pipeline: Document Extractor.

Primary extraction path uses Docling (OCR disabled, table extraction
disabled per spec) to convert PDF / DOCX / TXT into markdown, which is
then converted to raw text.

Docling's default pipeline downloads layout models on first use. If
Docling is unavailable (not installed, or no network access to fetch its
models), the extractor transparently falls back to PyMuPDF (for PDF),
python-docx (for DOCX) or a plain read (for TXT) so the system still
produces usable raw text end-to-end.
"""

from __future__ import annotations

import logging
from pathlib import Path

from src.models import Document


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class UnsupportedFileTypeError(ValueError):
    """Raised when a file extension is not one of the supported types."""


class DocumentExtractor:
    """Extracts raw text + page count from PDF / DOCX / TXT files."""

    def __init__(self, logger: logging.Logger, enable_ocr: bool = False,
                 enable_table_extraction: bool = False) -> None:
        self.logger = logger
        self.enable_ocr = enable_ocr
        self.enable_table_extraction = enable_table_extraction

    def extract(self, file_path: Path, output_dir: Path | None = None) -> Document:
        self.logger.info("Extracting document")
        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFileTypeError(f"Unsupported file type: {suffix}")

        from src.rag.multimodal_extractor import MultimodalExtractor
        
        # Initialize MultimodalExtractor. We configure it to enable OCR if requested
        # by the pipeline configuration, and always enable tables and images for RAG.
        extractor = MultimodalExtractor(
            enable_ocr=self.enable_ocr,
            enable_table_extraction=True,
            enable_image_extraction=True
        )
        
        raw_text, structured_doc, page_count = extractor.extract(file_path, output_dir)
        
        # Save structured_document.json to output_dir if provided
        if output_dir:
            structured_doc_path = output_dir / "structured_document.json"
            try:
                if hasattr(structured_doc, "model_dump_json"):
                    json_data = structured_doc.model_dump_json(indent=2)
                else:
                    json_data = structured_doc.json(indent=2)
                structured_doc_path.write_text(json_data, encoding="utf-8")
                self.logger.info("Saved structured_document.json to %s", structured_doc_path)
            except Exception as e:
                self.logger.error("Failed to save structured_document.json: %s", e)

        document = Document(
            name=file_path.stem,
            source_path=str(file_path),
            file_type=suffix.lstrip("."),
            raw_text=raw_text,
            page_count=page_count,
        )
        self.logger.info(
            "Extraction complete: %d characters, %d page(s)",
            len(raw_text), page_count,
        )
        return document

    # ------------------------------------------------------------------
    # Docling primary path
    # ------------------------------------------------------------------
    def _extract_with_docling(self, file_path: Path, suffix: str):
        try:
            from docling.document_converter import DocumentConverter
            from docling.datamodel.pipeline_options import PdfPipelineOptions
            from docling.document_converter import PdfFormatOption
            from docling.datamodel.base_models import InputFormat
        except ImportError:
            self.logger.warning("Docling not installed, skipping Docling extraction path")
            return None, 0

        try:
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = self.enable_ocr
            pipeline_options.do_table_structure = self.enable_table_extraction

            converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
                }
            )
            result = converter.convert(str(file_path))
            markdown = result.document.export_to_markdown()
            raw_text = self._markdown_to_text(markdown)
            page_count = len(getattr(result.document, "pages", []) or []) or 1
            return raw_text, page_count
        except Exception as exc:  # noqa: BLE001 - broad by design, we always fall back
            self.logger.warning("Docling extraction failed (%s); using fallback extractor", exc)
            return None, 0

    @staticmethod
    def _markdown_to_text(markdown: str) -> str:
        from src.utils import strip_markdown
        return strip_markdown(markdown)

    # ------------------------------------------------------------------
    # Fallback extraction path (no external model downloads required)
    # ------------------------------------------------------------------
    def _extract_fallback(self, file_path: Path, suffix: str):
        if suffix == ".pdf":
            return self._extract_pdf_fallback(file_path)
        if suffix == ".docx":
            return self._extract_docx_fallback(file_path)
        if suffix == ".txt":
            return self._extract_txt_fallback(file_path)
        return "", 0

    def _extract_pdf_fallback(self, file_path: Path):
        import fitz  # PyMuPDF

        text_parts = []
        with fitz.open(file_path) as pdf:
            page_count = pdf.page_count
            for page in pdf:
                text_parts.append(page.get_text("text"))
        return "\n\n".join(text_parts), page_count

    def _extract_docx_fallback(self, file_path: Path):
        import docx

        doc = docx.Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs]
        text = "\n".join(paragraphs)
        # DOCX has no fixed page concept; approximate one "page".
        return text, 1

    def _extract_txt_fallback(self, file_path: Path):
        text = file_path.read_text(encoding="utf-8", errors="ignore")
        return text, 1
